import pk.ioutil
from docx import Document
from docx.shared import Cm
from pk import decoratorsFunc
import pk.sfz


class DocxUtil():

    def __init__(self):
        self.sms_templet = Document(r".\templet\房产面积测算说明书.docx")
        self.chjssms_templet = Document(r".\templet\农村宅基地使用权及房屋所有权测绘技术说明.docx")
        self.sdckb_templet = Document(r".\templet\不动产实地查看记录表.docx")
        self.execldata = pk.ioutil.IoUtil()

    @decoratorsFunc.getexceptionreturn
    def checksfz(self):
        sfz_list = self.execldata.list_excel("身份证")
        sfz_list = [str(i) for i in sfz_list]
        execldata_sfz = pk.sfz.Sfz(sfz_list)
        check_sfz = execldata_sfz.check_list()
        return check_sfz

    #   修改段落内容
    def makevalue(self, doc_templet: Document, paragraph: int, runs: int, value, string):
        para = doc_templet.paragraphs
        para[paragraph].runs[runs].text = para[paragraph].runs[runs].text.replace(value, string)

    #   修改doc文档插入表格的内容
    def maketablevalue(self, doc_templet: Document, table: int, columu: int, row: int, paragraph: int, runs: int,
                       value: str, string: str):
        doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].runs[runs].text = \
            doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].runs[
                runs].text.replace(value, string)

    def maketableparagraphs(self, doc_templet: Document, table: int, columu: int, row: int, paragraph: int,
                            value: str, string: str):
        doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text = \
            doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text.replace(value, string)

    #   向docx文档插入的表格添加内容
    def addtablevalue(self, doc_templet: Document, table: int, columu: int, row: int, paragraph: int, value):  # 向
        doc_templet.tables[table].cell(columu, row).paragraphs[paragraph].text = value

    def getsms(self, save_path):
        pass

    #   房产面积测算说明书
    # @decoratorsFunc.getexceptionreturn
    def getsmss(self, save_path):
        n = 0
        for i in range(0, len(self.execldata.list_excel("宗地代码"))):
            self.sms_templet = Document(r"templet\房产面积测算说明书.docx")
            value = self.execldata.row_excel(i)

            #   转换日期数据格式
            value[5] = value[5][:4] + "年" + value[5][5:7] + "月" + value[5][8:10] + "日"
            value[4] = value[4][:4] + "年" + value[4][5:7] + "月" + value[4][8:10] + "日"

            self.makevalue(self.sms_templet, 0, 0, "F00010001", value[1])  # 不动产单元号：F00010001
            self.makevalue(self.sms_templet, 11, 1, "xxx", value[2])  # 委托测绘的房屋名称：xxx
            self.makevalue(self.sms_templet, 12, 1, "xxx", value[26])  # 委托测绘的房屋地址：xxx
            self.makevalue(self.sms_templet, 20, 0, "2020年11月23日", value[4])
            self.makevalue(self.sms_templet, 49, 3, "xx", value[2])  # 规划批准项目名称：  xx
            self.makevalue(self.sms_templet, 50, 3, "xx", value[2])  # 本次申请测绘的楼盘（组团）名称：  xx
            self.makevalue(self.sms_templet, 55, 1, "cc", str(value[22]))  # 层数
            self.makevalue(self.sms_templet, 55, 3, "dd", str(value[22]))  # 层数
            self.makevalue(self.sms_templet, 55, 5, "aa", str(value[17]))  # 总建筑面积为：aa平方米
            self.makevalue(self.sms_templet, 55, 7, "bb", str(value[16]))  # 建筑占地面积为：bb平方米
            self.makevalue(self.sms_templet, 57, 0, "xxx", value[33])  #
            if str(value[8]) != "nan":
                self.makevalue(self.sms_templet, 65, 1, "x", value[8])  # 根据委托方提供乡村建设规划许可证及附件，x
            else:
                self.makevalue(self.sms_templet, 65, 1, "x", "")
            self.makevalue(self.sms_templet, 106, 2, "xxx", value[2])  # 建筑物名称:xxx
            self.makevalue(self.sms_templet, 109, 2, "xxx", value[2])  # 建筑物名称: xxx
            self.makevalue(self.sms_templet, 112, 2, "xxx", value[2])  # 建筑物名称: xxx
            self.makevalue(self.sms_templet, 116, 2, "xxx", value[4])  # 计算者：  xxx
            self.makevalue(self.sms_templet, 117, 1, "xxx", value[4])  # 制表日期：xxx
            self.makevalue(self.sms_templet, 118, 0, "苏航", value[2])  # 苏航(近景)
            self.makevalue(self.sms_templet, 120, 0, "苏航", value[2])  # 苏航(远景)

            self.maketablevalue(self.sms_templet, 0, 0, 0, 1, 1, "2020年11月23日", value[4])  # 检查日期：   2020年11月23日
            self.maketablevalue(self.sms_templet, 0, 0, 0, 3, 0, "2020年11月25日", value[5])  # 检查日期：   2020年11月25日
            self.maketablevalue(self.sms_templet, 0, 1, 0, 1, 2, "2020年11月25日", value[5])  # 检查日期：   2020年11月25日

            self.maketablevalue(self.sms_templet, 1, 0, 3, 0, 0, "xxx", value[0])  # 宗地号
            self.maketablevalue(self.sms_templet, 1, 0, 6, 0, 0, "xx", value[2])  # 建筑物名称
            self.maketablevalue(self.sms_templet, 1, 1, 3, 0, 0, "xxx", value[7])  # 地址
            self.maketablevalue(self.sms_templet, 1, 4, 3, 0, 0, "xxx", str(value[16]))  # 基底面积
            self.maketablevalue(self.sms_templet, 1, 5, 3, 0, 0, "xxx", value[17])  # 总建筑面积
            self.maketablevalue(self.sms_templet, 1, 5, 6, 0, 0, "xxx", value[22])  # 地面以上层数
            self.maketablevalue(self.sms_templet, 1, 6, 3, 0, 0, "xxx", value[17])  # 地面以上面积
            # self.maketablevalue(1, 9, 3, 0, 0, "xxx", value[17])  # 共用建筑面积总计

            i = 18
            fcmj = []
            while str(value[i]) != "nan":
                fcmj.append(str(value[i]))
                i += 1
            for i in range(0, len(fcmj)):
                self.addtablevalue(self.sms_templet, 2, 1 + i, 0, 0, "第1幢 第" + str(i + 1) + "层")  # 第一列
                self.addtablevalue(self.sms_templet, 2, 1 + i, 1, 0, fcmj[i])  # 第二列
                self.addtablevalue(self.sms_templet, 2, 1 + i, 2, 0, fcmj[i] + "*1")  # 第三列
            self.addtablevalue(self.sms_templet, 2, 29, 1, 0, value[17])
            for i in range(0, len(fcmj)):
                self.addtablevalue(self.sms_templet, 4, 1 + i, 0, 0, r"第1")  # 第一列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 1, 0, r"/")  # 第二列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 2, 0, "第" + str(i + 1))  # 第三列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 3, 0, "1")  # 第四列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 4, 0, fcmj[i])  # 第五列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 5, 0, "0")  # 第六列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 6, 0, "0")  # 第七列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 7, 0, fcmj[i])  # 第八列
                self.addtablevalue(self.sms_templet, 4, 1 + i, 8, 0, "1")  # 第九列
            self.addtablevalue(self.sms_templet, 4, 24, 4, 0, value[17])
            self.addtablevalue(self.sms_templet, 4, 24, 7, 0, value[17])
            self.setjpg(r"%s\%s%s\近景.jpg" % (save_path, value[0], value[2]),
                        r"%s\%s%s\远景.jpg" % (save_path, value[0], value[2]))

            self.sms_templet.save(r"%s\%s%s\%s房产面积测算说明书.doc" % (save_path, value[0], value[2], value[0]))
            n += 1
        results = "房产面积测算说明书生成： %s个" % n
        return results

    #   不动产实地查看记录表
    @decoratorsFunc.getexceptionreturn
    def getsdckb(self, save_path):
        zddm = self.execldata.list_excel("宗地代码")
        zl = self.execldata.list_excel("坐落")
        xm = self.execldata.list_excel("姓名")
        dcrq = self.execldata.list_excel("调查日期")
        n = 0
        while n < len(xm):
            dcrq[n] = "%s年%s月%s日" % (str(dcrq[n])[:4], str(dcrq[n])[5:7], str(dcrq[n])[8:10])
            self.maketableparagraphs(self.sdckb_templet, 0, 0, 3, 0, "[xm]", xm[n])
            self.maketableparagraphs(self.sdckb_templet, 0, 1, 1, 0, "[zl]", zl[n])
            self.maketableparagraphs(self.sdckb_templet, 0, 4, 1, 1, "[dcrq]", dcrq[n])
            self.sdckb_templet.save(r"%s\%s%s\%s不动产实地查看记录表.doc" % (save_path, zddm[n], xm[n], zddm[n]))
            n += 1

        results = "不动产实地查看记录表生成：%s个" % n
        return results

    #   农村宅基地使用权及房屋所有权测绘技术说明
    @decoratorsFunc.getexceptionreturn
    def getchjssms(self, save_path):
        zddm = self.execldata.list_excel("宗地代码")
        xm = self.execldata.list_excel("姓名")
        dcrq = self.execldata.list_excel("调查日期")

        zdmj = self.execldata.list_excel("宗地面积")
        jzmj = self.execldata.list_excel("建筑总面积")

        i = 0
        while i < len(xm):
            self.chjssms_templet = Document(r"templet\农村宅基地使用权及房屋所有权测绘技术说明.docx")
            dcrq[i] = "%s年%s月%s日" % (str(dcrq[i])[:4], str(dcrq[i])[5:7], str(dcrq[i])[8:10])
            self.makevalue(self.chjssms_templet, 1, 0, "aaa", xm[i])
            self.makevalue(self.chjssms_templet, 1, 2, "bbb", dcrq[i])
            self.makevalue(self.chjssms_templet, 1, 6, "ccc", xm[i])
            self.makevalue(self.chjssms_templet, 7, 1, "ddd", str(zdmj[i]))
            self.makevalue(self.chjssms_templet, 7, 3, "eee", str(jzmj[i]))
            self.makevalue(self.chjssms_templet, 11, 0, "xxx", dcrq[i])
            self.chjssms_templet.save(r"%s\%s%s\农村宅基地使用权及房屋所有权测绘技术说明（%s）.doc" % (save_path, zddm[i], xm[i], xm[i]))

            i += 1
        results = "农村宅基地使用权及房屋所有权测绘技术说明生成：%s个" % i
        return results

    def setjpg(self, path_JJ, path_YJ):
        self.sms_templet.paragraphs[119].add_run().add_picture(path_JJ, width=Cm(14))
        self.sms_templet.add_picture(path_YJ, width=Cm(14))


if __name__ == '__main__':
    mydo = DocxUtil()
    string1 = mydo.sdckb_templet.tables[0].cell(4, 1).paragraphs[1].text
    string2 = mydo.sms_templet.tables[0].cell(0, 1).paragraphs[1].text
    string3 = mydo.sms_templet.paragraphs[12].text
    string4 = mydo.sms_templet.paragraphs[12].runs[1].text

    print(string1)
    print(string2)
    print(string3)
    print(string4)
